import os
import shutil
import hashlib
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import json
import httpx
import asyncio

# LangChain imports
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_classic.chains import RetrievalQA
try:
    from langchain_core.prompts import PromptTemplate
    from langchain_core.documents import Document
except ImportError:
    from langchain.prompts import PromptTemplate
    from langchain.schema import Document

# Import para DOCX
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

app = FastAPI(title="DevOpsAI Backend")

# Configurações
PERSIST_DIRECTORY = "/app/chroma_db"
OLLAMA_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
DEFAULT_MODEL = os.getenv("DEFAULT_MODEL", "deepseek-r1:1.5b")

embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

def get_vectorstore():
    """Retorna o vectorstore ChromaDB persistente."""
    return Chroma(
        persist_directory=PERSIST_DIRECTORY, 
        embedding_function=embeddings
    )

def get_file_hash(file_path: str) -> str:
    """Calcula o hash MD5 de um arquivo."""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def load_document(file_path: str, file_extension: str) -> List[Document]:
    """Carrega documentos de diferentes formatos."""
    extension = file_extension.lower()
    
    if extension == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif extension == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    elif extension in [".md", ".markdown"]:
        # Carrega Markdown como texto simples
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    elif extension == ".docx":
        if DocxDocument is None:
            raise HTTPException(
                status_code=400, 
                detail="Suporte a DOCX requer python-docx. Instale com: pip install python-docx"
            )
        # Carrega DOCX manualmente
        doc = DocxDocument(file_path)
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        # Também extrai texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        full_text = "\n".join(text_content)
        if not full_text.strip():
            raise HTTPException(status_code=400, detail="Arquivo DOCX está vazio ou não pôde ser lido")
        
        return [Document(page_content=full_text, metadata={"source": file_path})]
    else:
        raise HTTPException(
            status_code=400, 
            detail=f"Formato não suportado: {extension}. Formatos suportados: PDF, TXT, MD, DOCX"
        )

def extract_keywords(question: str) -> List[str]:
    """Extrai palavras-chave significativas da pergunta."""
    words = question.lower().split()
    stop_words = {'o', 'a', 'os', 'as', 'um', 'uma', 'de', 'da', 'do', 'em', 'na', 'no', 'sobre', 'para', 'com', 'que', 'me', 'você'}
    keywords = [w for w in words if len(w) > 3 and w not in stop_words]
    return keywords

def find_relevant_documents(vectorstore, question: str, keywords: List[str], max_docs: int = 20) -> List[Document]:
    """Busca documentos relevantes usando busca semântica e filtro por palavras-chave."""
    relevant_docs = []
    
    try:
        semantic_docs = vectorstore.similarity_search(question, k=max_docs)
    except Exception:
        try:
            retriever = vectorstore.as_retriever(search_kwargs={"k": max_docs})
            semantic_docs = retriever.invoke(question)
        except Exception:
            semantic_docs = []
    
    question_lower = question.lower()
    for doc in semantic_docs:
        content_lower = doc.page_content.lower()
        if any(keyword in content_lower for keyword in keywords) or any(word in content_lower for word in question_lower.split() if len(word) > 3):
            relevant_docs.append(doc)
    
    if not relevant_docs:
        all_data = vectorstore.get()
        if all_data.get('documents'):
            for i, content in enumerate(all_data['documents']):
                content_lower = content.lower()
                if any(keyword in content_lower for keyword in keywords):
                    metadata = all_data['metadatas'][i] if all_data.get('metadatas') and i < len(all_data['metadatas']) else {}
                    doc = Document(page_content=content, metadata=metadata)
                    relevant_docs.append(doc)
                    if len(relevant_docs) >= max_docs:
                        break
    
    return relevant_docs

def get_llm(model_name: str = None, temperature: float = 0.7, top_p: float = 0.9):
    """Cria uma instância do LLM com configurações customizadas."""
    model = model_name or DEFAULT_MODEL
    return ChatOllama(
        base_url=OLLAMA_URL, 
        model=model,
        temperature=temperature,
        top_p=top_p,
        timeout=120.0  # Timeout de 2 minutos
    )

async def verify_model_exists(model_name: str) -> bool:
    """Verifica se um modelo está instalado no Ollama."""
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.post(
                f"{OLLAMA_URL}/api/show",
                json={"name": model_name}
            )
            return response.status_code == 200
    except:
        return False

# Modelos Pydantic
class QueryRequest(BaseModel):
    question: str
    model: Optional[str] = None
    temperature: Optional[float] = Field(default=0.7, ge=0.0, le=2.0)
    top_p: Optional[float] = Field(default=0.9, ge=0.0, le=1.0)
    stream: Optional[bool] = False

class DocumentMetadata(BaseModel):
    filename: str
    chunks: int
    file_hash: Optional[str] = None
    custom_metadata: Optional[Dict[str, Any]] = None

class DeleteDocumentRequest(BaseModel):
    filename: str

class UpdateDocumentRequest(BaseModel):
    filename: str
    custom_metadata: Dict[str, Any]

@app.get("/health")
async def health():
    """Endpoint de healthcheck."""
    return {"status": "healthy", "ollama_url": OLLAMA_URL, "default_model": DEFAULT_MODEL}

# Modelos conhecidos disponíveis para download
KNOWN_MODELS = {
    "deepseek-r1:1.5b": {
        "name": "DeepSeek R1 1.5B",
        "size": "~1.1GB",
        "description": "Modelo rápido e eficiente, ideal para respostas rápidas",
        "tags": ["rápido", "leve", "eficiente"]
    },
    "deepseek-r1:8b": {
        "name": "DeepSeek R1 8B",
        "size": "~4.7GB",
        "description": "Melhor qualidade, requer mais RAM",
        "tags": ["qualidade", "médio"]
    },
    "deepseek": {
        "name": "DeepSeek Chat",
        "size": "~4.7GB",
        "description": "Versão padrão do DeepSeek Chat",
        "tags": ["padrão", "chat"]
    },
    "deepseek-coder": {
        "name": "DeepSeek Coder",
        "size": "~4.7GB",
        "description": "Especializado em código e programação",
        "tags": ["código", "programação"]
    },
    "llama3.2": {
        "name": "Llama 3.2",
        "size": "~2GB",
        "description": "Modelo balanceado da Meta",
        "tags": ["balanceado", "geral"]
    },
    "llama3": {
        "name": "Llama 3",
        "size": "~4.7GB",
        "description": "Versão completa do Llama 3",
        "tags": ["completo", "geral"]
    },
    "llama3.1": {
        "name": "Llama 3.1",
        "size": "~4.7GB",
        "description": "Versão 3.1 do Llama",
        "tags": ["geral"]
    },
    "mistral": {
        "name": "Mistral",
        "size": "~4.1GB",
        "description": "Modelo Mistral de alta qualidade",
        "tags": ["qualidade", "geral"]
    },
    "phi3:mini": {
        "name": "Phi-3 Mini",
        "size": "~2.3GB",
        "description": "Modelo leve e rápido da Microsoft",
        "tags": ["leve", "rápido"]
    },
    "gemma2:2b": {
        "name": "Gemma 2 2B",
        "size": "~1.4GB",
        "description": "Modelo leve do Google",
        "tags": ["leve", "google"]
    },
    "qwen2.5:7b": {
        "name": "Qwen 2.5 7B",
        "size": "~4.4GB",
        "description": "Modelo de alta qualidade",
        "tags": ["qualidade"]
    }
}

async def get_ollama_models() -> List[Dict[str, Any]]:
    """Busca modelos instalados no Ollama."""
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(f"{OLLAMA_URL}/api/tags")
            if response.status_code == 200:
                data = response.json()
                models = []
                for model in data.get("models", []):
                    model_name = model.get("name", "")
                    models.append({
                        "name": model_name,
                        "size": model.get("size", 0),
                        "modified_at": model.get("modified_at", ""),
                        "digest": model.get("digest", "")
                    })
                return models
            return []
    except Exception as e:
        print(f"Erro ao buscar modelos do Ollama: {e}")
        return []

@app.get("/models/installed")
async def list_installed_models():
    """Lista modelos instalados no Ollama."""
    try:
        models = await get_ollama_models()
        return {
            "installed_models": models,
            "count": len(models)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao listar modelos: {str(e)}")

@app.get("/models/available")
async def list_available_models():
    """Lista modelos disponíveis para download."""
    installed_models = await get_ollama_models()
    installed_names = {model["name"] for model in installed_models}
    
    available = []
    for model_id, info in KNOWN_MODELS.items():
        available.append({
            "id": model_id,
            "name": info["name"],
            "size": info["size"],
            "description": info["description"],
            "tags": info["tags"],
            "installed": model_id in installed_names or any(model_id in name for name in installed_names)
        })
    
    return {
        "available_models": available,
        "count": len(available)
    }

@app.get("/models")
async def list_models():
    """Endpoint combinado que retorna modelos instalados e disponíveis."""
    installed = await get_ollama_models()
    installed_names = {model["name"] for model in installed}
    
    available = []
    for model_id, info in KNOWN_MODELS.items():
        available.append({
            "id": model_id,
            "name": info["name"],
            "size": info["size"],
            "description": info["description"],
            "tags": info["tags"],
            "installed": model_id in installed_names or any(model_id in name for name in installed_names)
        })
    
    return {
        "installed_models": installed,
        "available_models": available,
        "default_model": DEFAULT_MODEL,
        "ollama_url": OLLAMA_URL
    }

class PullModelRequest(BaseModel):
    model: str

async def stream_model_pull(model_name: str):
    """Stream do download de modelo do Ollama."""
    try:
        # Timeout maior para modelos grandes (30 minutos)
        timeout = httpx.Timeout(1800.0, connect=30.0)
        async with httpx.AsyncClient(timeout=timeout) as client:
            async with client.stream(
                "POST",
                f"{OLLAMA_URL}/api/pull",
                json={"name": model_name},
                timeout=timeout
            ) as response:
                # Verifica se a resposta HTTP está OK
                if response.status_code != 200:
                    error_text = await response.aread()
                    error_data = {
                        "error": f"Erro HTTP {response.status_code}: {error_text.decode('utf-8', errors='ignore')}",
                        "status": "error"
                    }
                    yield f"data: {json.dumps(error_data)}\n\n"
                    return
                
                async for line in response.aiter_lines():
                    if line:
                        try:
                            data = json.loads(line)
                            
                            # Verifica se há erro na resposta do Ollama
                            if "error" in data:
                                error_data = {
                                    "error": data.get("error", "Erro desconhecido do Ollama"),
                                    "status": "error"
                                }
                                yield f"data: {json.dumps(error_data)}\n\n"
                                return
                            
                            # Propaga mensagens de progresso e sucesso
                            yield f"data: {json.dumps(data)}\n\n"
                            
                            # Se o status for "success", finaliza
                            if data.get("status") == "success":
                                return
                                
                        except json.JSONDecodeError:
                            # Ignora linhas que não são JSON válido
                            continue
                        except Exception as e:
                            error_data = {
                                "error": f"Erro ao processar resposta: {str(e)}",
                                "status": "error"
                            }
                            yield f"data: {json.dumps(error_data)}\n\n"
                            return
                            
    except httpx.TimeoutException as e:
        error_data = {
            "error": f"Timeout ao baixar modelo. O download pode ter sido interrompido. Tente novamente.",
            "status": "error",
            "details": str(e)
        }
        yield f"data: {json.dumps(error_data)}\n\n"
    except httpx.ConnectError as e:
        error_data = {
            "error": f"Não foi possível conectar com Ollama. Verifique se o serviço está rodando.",
            "status": "error",
            "details": str(e)
        }
        yield f"data: {json.dumps(error_data)}\n\n"
    except Exception as e:
        error_data = {
            "error": f"Erro ao baixar modelo: {str(e)}",
            "status": "error"
        }
        yield f"data: {json.dumps(error_data)}\n\n"

@app.post("/models/pull")
async def pull_model(request: PullModelRequest):
    """Baixa um modelo do Ollama com streaming de progresso."""
    # Valida se o modelo está na lista conhecida
    if request.model not in KNOWN_MODELS:
        # Permite modelos customizados também
        pass
    
    return StreamingResponse(
        stream_model_pull(request.model),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )

@app.delete("/models/{model_name}")
async def delete_model(model_name: str):
    """Remove um modelo do Ollama."""
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.delete(
                f"{OLLAMA_URL}/api/delete",
                json={"name": model_name}
            )
            if response.status_code == 200:
                return {
                    "message": f"Modelo '{model_name}' removido com sucesso",
                    "status": "success"
                }
            else:
                error_data = response.json() if response.content else {}
                raise HTTPException(
                    status_code=response.status_code,
                    detail=error_data.get("error", "Erro ao remover modelo")
                )
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Timeout ao conectar com Ollama")
    except httpx.ConnectError:
        raise HTTPException(status_code=503, detail="Não foi possível conectar com Ollama")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao remover modelo: {str(e)}")

@app.get("/models/{model_name}/info")
async def get_model_info(model_name: str):
    """Retorna informações sobre um modelo específico."""
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.post(
                f"{OLLAMA_URL}/api/show",
                json={"name": model_name}
            )
            if response.status_code == 200:
                return response.json()
            else:
                raise HTTPException(
                    status_code=response.status_code,
                    detail="Modelo não encontrado"
                )
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Timeout ao conectar com Ollama")
    except httpx.ConnectError:
        raise HTTPException(status_code=503, detail="Não foi possível conectar com Ollama")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao buscar informações: {str(e)}")

@app.get("/documents", response_model=List[Dict[str, Any]])
async def list_documents():
    """Lista documentos indexados com metadados."""
    try:
        vectorstore = get_vectorstore()
        data = vectorstore.get()
        
        if not data or not data['metadatas']:
            return []

        # Agrupa por arquivo
        file_info = {}
        for i, metadata in enumerate(data['metadatas']):
            if metadata and 'source' in metadata:
                filename = os.path.basename(metadata['source'])
                if filename.startswith("temp_"):
                    filename = filename[5:]
                
                if filename not in file_info:
                    file_info[filename] = {
                        "filename": filename,
                        "chunks": 0,
                        "file_hash": metadata.get("file_hash"),
                        "custom_metadata": metadata.get("custom_metadata", {})
                    }
                file_info[filename]["chunks"] += 1
        
        return list(file_info.values())
    except Exception as e:
        return []

@app.get("/documents/{filename}")
async def get_document_info(filename: str):
    """Retorna informações detalhadas sobre um documento específico."""
    try:
        vectorstore = get_vectorstore()
        data = vectorstore.get()
        
        if not data or not data['metadatas']:
            raise HTTPException(status_code=404, detail="Documento não encontrado")
        
        chunks = []
        file_hash = None
        custom_metadata = {}
        
        for i, metadata in enumerate(data['metadatas']):
            if metadata and 'source' in metadata:
                meta_filename = os.path.basename(metadata['source'])
                if meta_filename.startswith("temp_"):
                    meta_filename = meta_filename[5:]
                
                if meta_filename == filename:
                    chunks.append({
                        "id": data['ids'][i] if i < len(data['ids']) else None,
                        "content_preview": data['documents'][i][:200] + "..." if i < len(data['documents']) else "",
                        "metadata": metadata
                    })
                    if not file_hash and metadata.get("file_hash"):
                        file_hash = metadata["file_hash"]
                    if metadata.get("custom_metadata"):
                        custom_metadata.update(metadata["custom_metadata"])
        
        if not chunks:
            raise HTTPException(status_code=404, detail="Documento não encontrado")
        
        return {
            "filename": filename,
            "chunks_count": len(chunks),
            "file_hash": file_hash,
            "custom_metadata": custom_metadata,
            "chunks": chunks[:10]  # Limita a 10 chunks para não sobrecarregar
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/documents/{filename}")
async def delete_document(filename: str):
    """Remove um documento da base de conhecimento."""
    try:
        vectorstore = get_vectorstore()
        data = vectorstore.get()
        
        if not data or not data.get('ids'):
            raise HTTPException(status_code=404, detail="Nenhum documento encontrado na base")
        
        # Encontra todos os IDs relacionados ao arquivo
        ids_to_delete = []
        for i, metadata in enumerate(data['metadatas']):
            if metadata and 'source' in metadata:
                meta_filename = os.path.basename(metadata['source'])
                if meta_filename.startswith("temp_"):
                    meta_filename = meta_filename[5:]
                
                if meta_filename == filename:
                    if i < len(data['ids']):
                        ids_to_delete.append(data['ids'][i])
        
        if not ids_to_delete:
            raise HTTPException(status_code=404, detail=f"Documento '{filename}' não encontrado")
        
        # Deleta os chunks
        vectorstore.delete(ids=ids_to_delete)
        
        return {
            "message": f"Documento '{filename}' removido com sucesso",
            "chunks_deleted": len(ids_to_delete)
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/documents/{filename}/metadata")
async def update_document_metadata(filename: str, request: UpdateDocumentRequest):
    """Atualiza metadados customizados de um documento."""
    try:
        vectorstore = get_vectorstore()
        data = vectorstore.get()
        
        if not data or not data.get('ids'):
            raise HTTPException(status_code=404, detail="Nenhum documento encontrado")
        
        updated_count = 0
        for i, metadata in enumerate(data['metadatas']):
            if metadata and 'source' in metadata:
                meta_filename = os.path.basename(metadata['source'])
                if meta_filename.startswith("temp_"):
                    meta_filename = meta_filename[5:]
                
                if meta_filename == filename:
                    # Atualiza metadados
                    new_metadata = metadata.copy()
                    if 'custom_metadata' not in new_metadata:
                        new_metadata['custom_metadata'] = {}
                    new_metadata['custom_metadata'].update(request.custom_metadata)
                    
                    # Atualiza no vectorstore
                    if i < len(data['ids']):
                        vectorstore.update_document(
                            id=data['ids'][i],
                            document=Document(
                                page_content=data['documents'][i] if i < len(data['documents']) else "",
                                metadata=new_metadata
                            )
                        )
                        updated_count += 1
        
        if updated_count == 0:
            raise HTTPException(status_code=404, detail=f"Documento '{filename}' não encontrado")
        
        return {
            "message": f"Metadados do documento '{filename}' atualizados",
            "chunks_updated": updated_count
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    custom_metadata: Optional[str] = None
):
    """Upload e processamento de arquivos (PDF, TXT, MD, DOCX)."""
    # Valida extensão
    file_extension = os.path.splitext(file.filename)[1]
    if file_extension.lower() not in [".pdf", ".txt", ".md", ".markdown", ".docx"]:
        raise HTTPException(
            status_code=400,
            detail=f"Formato não suportado. Use: PDF, TXT, MD, ou DOCX"
        )
    
    # Valida tamanho (máximo 50MB)
    file.file.seek(0, 2)  # Vai para o final
    file_size = file.file.tell()
    file.file.seek(0)  # Volta para o início
    
    if file_size > 50 * 1024 * 1024:  # 50MB
        raise HTTPException(status_code=400, detail="Arquivo muito grande. Máximo: 50MB")
    
    # Salva temporariamente
    temp_filename = f"temp_{file.filename}"
    with open(temp_filename, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    try:
        # Calcula hash do arquivo
        file_hash = get_file_hash(temp_filename)
        
        # Verifica se já existe (opcional - pode ser removido se quiser permitir duplicatas)
        vectorstore = get_vectorstore()
        data = vectorstore.get()
        if data and data.get('metadatas'):
            for metadata in data['metadatas']:
                if metadata and metadata.get('file_hash') == file_hash:
                    return {
                        "message": "Arquivo já existe na base de conhecimento",
                        "chunks": 0,
                        "file_hash": file_hash,
                        "duplicate": True
                    }
        
        # Carrega documento
        documents = load_document(temp_filename, file_extension)
        
        # Processa metadados customizados
        custom_meta_dict = {}
        if custom_metadata:
            try:
                custom_meta_dict = json.loads(custom_metadata)
            except:
                pass
        
        # Adiciona metadados aos documentos
        for doc in documents:
            doc.metadata["file_hash"] = file_hash
            if custom_meta_dict:
                doc.metadata["custom_metadata"] = custom_meta_dict
        
        # Divide em chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_documents(documents)
        
        # Adiciona ao vectorstore
        vectorstore.add_documents(chunks)
        
        return {
            "message": f"Arquivo '{file.filename}' processado com sucesso!",
            "chunks": len(chunks),
            "file_hash": file_hash,
            "format": file_extension
        }
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao processar arquivo: {str(e)}")
    finally:
        if os.path.exists(temp_filename):
            os.remove(temp_filename)

async def generate_streaming_response(question: str, context_text: str, llm, prompt_template: str):
    """Gera resposta em streaming."""
    PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    formatted_prompt = PROMPT.format(context=context_text, question=question)
    
    buffer = ""
    try:
        # Tenta usar astream se disponível
        if hasattr(llm, 'astream'):
            async for chunk in llm.astream(formatted_prompt):
                content = chunk.content if hasattr(chunk, 'content') else str(chunk)
                buffer += content
                yield f"data: {json.dumps({'chunk': content, 'done': False})}\n\n"
        else:
            # Fallback: usa stream síncrono em thread separada
            import asyncio
            loop = asyncio.get_event_loop()
            for chunk in await loop.run_in_executor(None, lambda: llm.stream(formatted_prompt)):
                content = chunk.content if hasattr(chunk, 'content') else str(chunk)
                buffer += content
                yield f"data: {json.dumps({'chunk': content, 'done': False})}\n\n"
    except Exception as e:
        error_msg = str(e)
        
        # Trata erros específicos do Ollama
        if "EOF" in error_msg or "status code: 500" in error_msg:
            error_message = "Erro ao comunicar com o Ollama. O modelo pode estar com problemas ou o contexto pode ser muito grande."
        elif "not found" in error_msg.lower() or "404" in error_msg:
            error_message = "Modelo não encontrado. Por favor, baixe o modelo primeiro."
        elif "timeout" in error_msg.lower():
            error_message = "Timeout ao processar. Tente novamente ou use um modelo menor."
        elif "Connection" in error_msg or "connect" in error_msg.lower():
            error_message = "Erro de conexão com o Ollama. Verifique se o serviço está rodando."
        else:
            error_message = f"Erro ao processar: {error_msg}"
        
        # Retorna erro em formato streaming
        yield f"data: {json.dumps({'chunk': '', 'error': error_message, 'done': True})}\n\n"
        return
    
    yield f"data: {json.dumps({'chunk': '', 'done': True, 'full_response': buffer})}\n\n"

@app.post("/chat")
async def chat(request: QueryRequest):
    """Endpoint de chat com suporte a streaming e configurações customizadas."""
    try:
        vectorstore = get_vectorstore()
        data = vectorstore.get()
        
        if not data or not data.get('ids') or len(data['ids']) == 0:
            return {
                "answer": "Não há documentos na base de conhecimento ainda. Por favor, faça upload de arquivos primeiro usando a opção na barra lateral."
            }
        
        # Extrai palavras-chave
        keywords = extract_keywords(request.question)
        
        # Busca documentos relevantes
        relevant_docs = find_relevant_documents(vectorstore, request.question, keywords, max_docs=20)
        
        if not relevant_docs:
            return {
                "answer": "Não encontrei informações relevantes na base de conhecimento para responder sua pergunta. Tente reformular a pergunta ou verifique se os documentos contêm essa informação."
            }
        
        # Limita documentos e tamanho do contexto
        relevant_docs = relevant_docs[:15]
        
        # Limita o tamanho do contexto para evitar problemas com modelos pequenos
        max_context_length = 8000  # Aproximadamente 8000 caracteres
        context_parts = []
        current_length = 0
        
        for doc in relevant_docs:
            doc_content = doc.page_content
            if current_length + len(doc_content) > max_context_length:
                # Adiciona apenas o que couber
                remaining = max_context_length - current_length
                if remaining > 100:  # Só adiciona se sobrar espaço significativo
                    context_parts.append(doc_content[:remaining])
                break
            context_parts.append(doc_content)
            current_length += len(doc_content)
        
        context_text = "\n\n".join(context_parts)
        
        # Determina o modelo a usar
        model_to_use = request.model or DEFAULT_MODEL
        
        # Verifica se o modelo existe
        model_exists = await verify_model_exists(model_to_use)
        if not model_exists:
            error_msg = f"❌ O modelo '{model_to_use}' não está instalado no Ollama. Por favor, baixe o modelo primeiro usando a seção 'Gerenciar Modelos' na barra lateral."
            
            # Se for streaming, retorna erro em formato streaming
            if request.stream:
                async def error_stream():
                    yield f"data: {json.dumps({'chunk': '', 'error': error_msg, 'done': True})}\n\n"
                
                return StreamingResponse(
                    error_stream(),
                    media_type="text/event-stream",
                    headers={
                        "Cache-Control": "no-cache",
                        "Connection": "keep-alive",
                        "X-Accel-Buffering": "no"
                    }
                )
            else:
                return {
                    "answer": error_msg
                }
        
        # Cria LLM com configurações customizadas
        llm = get_llm(
            model_name=model_to_use,
            temperature=request.temperature,
            top_p=request.top_p
        )
        
        prompt_template = """Você é um assistente especializado em responder perguntas baseado APENAS no contexto fornecido da base de conhecimento interna.

INSTRUÇÕES CRÍTICAS:
- Use EXCLUSIVAMENTE as informações do contexto abaixo
- Se a informação não estiver no contexto, diga claramente: "Não encontrei essa informação na base de conhecimento"
- NÃO invente informações
- NÃO use conhecimento externo ou informações gerais que não estejam no contexto
- NÃO mencione tecnologias ou ferramentas que não estejam explicitamente no contexto fornecido
- Foque APENAS no que está documentado no contexto

Contexto da Base de Conhecimento:
{context}

Pergunta do usuário: {question}

Resposta (baseada EXCLUSIVAMENTE no contexto acima):"""
        
        # Se streaming está habilitado
        if request.stream:
            return StreamingResponse(
                generate_streaming_response(request.question, context_text, llm, prompt_template),
                media_type="text/event-stream",
                headers={
                    "Cache-Control": "no-cache",
                    "Connection": "keep-alive",
                    "X-Accel-Buffering": "no"
                }
            )
        
        # Resposta não-streaming
        PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
        formatted_prompt = PROMPT.format(context=context_text, question=request.question)
        
        try:
            llm_response = llm.invoke(formatted_prompt)
            answer = llm_response.content if hasattr(llm_response, 'content') else str(llm_response)
            
            return {
                "answer": answer,
                "sources_count": len(relevant_docs),
                "model": model_to_use
            }
        except Exception as llm_error:
            error_msg = str(llm_error)
            
            # Trata erros específicos do Ollama
            if "EOF" in error_msg or "status code: 500" in error_msg:
                return {
                    "answer": f"❌ Erro ao comunicar com o Ollama. O modelo '{model_to_use}' pode estar com problemas ou o contexto pode ser muito grande. Tente:\n\n1. Verificar se o modelo está instalado corretamente\n2. Reduzir o número de documentos na base de conhecimento\n3. Tentar novamente em alguns instantes"
                }
            elif "not found" in error_msg.lower() or "404" in error_msg:
                return {
                    "answer": f"❌ O modelo '{model_to_use}' não foi encontrado. Por favor, baixe o modelo primeiro usando a seção 'Gerenciar Modelos' na barra lateral."
                }
            elif "timeout" in error_msg.lower():
                return {
                    "answer": f"⏱️ Timeout ao processar a resposta. O modelo pode estar demorando muito. Tente:\n\n1. Reduzir o tamanho do contexto\n2. Usar um modelo mais rápido\n3. Tentar novamente"
                }
            else:
                return {
                    "answer": f"❌ Erro ao processar a resposta: {error_msg}\n\nPor favor, tente novamente ou verifique se o modelo está funcionando corretamente."
                }
        
    except Exception as e:
        import traceback
        error_msg = str(e)
        
        # Mensagens de erro mais amigáveis
        if "Connection" in error_msg or "connect" in error_msg.lower():
            return {
                "answer": "❌ Erro de conexão: Não foi possível conectar com o Ollama. Verifique se o serviço está rodando."
            }
        elif "timeout" in error_msg.lower():
            return {
                "answer": "⏱️ Timeout: A operação demorou muito. Tente novamente ou use um modelo menor."
            }
        else:
            return {
                "answer": f"❌ Erro ao processar: {error_msg}\n\nSe o problema persistir, verifique os logs do backend."
            }
